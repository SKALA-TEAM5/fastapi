from __future__ import annotations

"""ReportDraft JSON을 DOCX 샘플 보고서 레이아웃으로 렌더링합니다.

renderer는 DOCX 관련 책임을 ReportAgent 밖에 둡니다. 샘플 보고서가 있으면
복사해서 표를 채우고, 이슈 상세 섹션을 데이터 개수에 맞게 확장합니다.
"""

import argparse
import json
import shutil
from copy import deepcopy
from decimal import Decimal
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import _Cell, Table

try:
    from .schemas import (
        IssueDetailDraft,
        ReportDraft,
    )
except ImportError:  # 직접 스크립트로 실행할 때 사용하는 예비 경로입니다.
    from schemas import (  # 타입 검사에서는 상대 import 경로만 사용합니다.
        IssueDetailDraft,
        ReportDraft,
    )


DEFAULT_TEMPLATE_PATH = Path(
    r"C:\Users\Hyeon\Downloads\산업안전보건관리비_증빙검토_보고서_AR-2025-0312_v2.docx"
)
MIN_TABLE_COUNT = 11
EMPTY_LABEL = "-"


def render_report_draft_to_docx(
    draft: ReportDraft,
    output_path: str | Path,
    *,
    template_path: str | Path | None = DEFAULT_TEMPLATE_PATH,
) -> Path:
    """ReportDraft를 DOCX 파일로 렌더링합니다.

    report agent는 구조화 JSON까지만 책임집니다. 이 renderer는 그 JSON을
    샘플 DOCX 보고서 구조에 매핑하는 별도 경계입니다.
    """

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    template = Path(template_path) if template_path else None
    if template and template.exists():
        shutil.copyfile(template, output)
        document = Document(str(output))
        if len(document.tables) < MIN_TABLE_COUNT:
            document = _create_base_document()
    else:
        document = _create_base_document()

    _apply_document_defaults(document)
    _fill_report_tables(document, draft)
    document.save(str(output))
    return output


def _create_base_document() -> DocumentObject:
    document = Document()
    _apply_document_defaults(document)
    document.add_heading("산업안전보건관리비 집행 증빙 검토 결과 보고서", level=0)
    for title, rows, cols in [
        ("1. 표지 요약", 5, 2),
        ("2. 기본 정보 상세", 6, 4),
        ("3. 집행 금액 요약", 5, 4),
        ("4. 집행 항목 분류별 요약", 2, 4),
        ("5. 증빙 유형별 검증 현황", 2, 6),
        ("6. 세금 및 정산 결과", 2, 5),
        ("7. 항목별 적정성 검토 결과", 2, 5),
        ("8. 부적정/검토 필요 상세 1", 5, 2),
        ("9. 부적정/검토 필요 상세 2", 5, 2),
        ("10. 부적정/검토 필요 상세 3", 5, 2),
        ("11. 보완 필요 사항", 2, 5),
        ("12. 종합 의견", 1, 1),
    ]:
        document.add_paragraph(title)
        document.add_table(rows=rows, cols=cols)
    return document


def _apply_document_defaults(document: DocumentObject) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    style = next((item for item in document.styles if item.style_id == "Normal"), None)
    if style is None:
        style = document.styles["Normal"]
    style.font.name = "맑은 고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(9)


def _fill_report_tables(document: DocumentObject, draft: ReportDraft) -> None:
    """ReportDraft 섹션을 v2 샘플 보고서의 표 순서에 맞게 채웁니다."""

    tables = document.tables
    if len(tables) < MIN_TABLE_COUNT:
        raise ValueError(f"DOCX template must contain at least {MIN_TABLE_COUNT} tables.")

    _fill_cover(tables[0], draft)
    _fill_basic_info(tables[1], draft)
    _fill_amount_summary(tables[2], draft)
    _fill_category_summaries(tables[3], draft)
    _fill_evidence_summaries(tables[4], draft)
    _fill_tax_rows(tables[5], draft)
    _fill_item_reviews(tables[6], draft)
    issue_tables = _ensure_issue_detail_blocks(document, tables[7:10], len(draft.issue_details))
    _fill_issue_detail_blocks(issue_tables, draft.issue_details)
    _fill_issue_headings(document, draft.issue_details)
    _fill_supplement_actions(tables[10], draft)
    if len(tables) >= 12:
        _set_cell(tables[11].cell(0, 0), draft.overall_opinion or EMPTY_LABEL)
    else:
        _fill_overall_opinion_paragraphs(document, draft)


def _fill_cover(table: Table, draft: ReportDraft) -> None:
    rows = [
        ("현장명", draft.site_name),
        ("검토 대상 기간", draft.report_period_label),
        ("보고서 번호", draft.report_no),
        ("작성일", draft.written_date_label),
        ("작성 부서", _department_full_label(draft)),
    ]
    _ensure_rows(table, len(rows))
    for row_index, row in enumerate(rows):
        _set_row(table, row_index, row)


def _fill_basic_info(table: Table, draft: ReportDraft) -> None:
    info = draft.basic_info
    rows = [
        ("보고서 번호", draft.report_no, "검토 일자", draft.written_date_label),
        ("현장명", draft.site_name, "현장코드", info.get("현장코드", EMPTY_LABEL)),
        ("발주처", info.get("발주처", EMPTY_LABEL), "시공사", info.get("시공사", EMPTY_LABEL)),
        ("계약금액", info.get("계약금액", EMPTY_LABEL), "공사기간", info.get("공사기간", EMPTY_LABEL)),
        ("법정 계상액", info.get("법정 계상액", EMPTY_LABEL), "검토자", draft.reviewer_label),
        ("검토 대상 기간", draft.report_period_label, "검토 목적", info.get("검토 목적", EMPTY_LABEL)),
    ]
    _ensure_rows(table, len(rows))
    for row_index, row in enumerate(rows):
        _set_row(table, row_index, row)


def _fill_amount_summary(table: Table, draft: ReportDraft) -> None:
    rows = [("구분", "금액", "집행률", "비고")]
    rows.extend(
        (
            summary.label,
            _money(summary.amount),
            summary.ratio_label,
            summary.count_label,
        )
        for summary in draft.amount_summary
    )
    _fill_grid(table, rows)


def _fill_category_summaries(table: Table, draft: ReportDraft) -> None:
    rows = [("집행 항목", "집행액 (원)", "건수", "비고")]
    rows.extend(
        (
            summary.category_name,
            _money(summary.amount),
            _count(summary.count),
            summary.note,
        )
        for summary in draft.category_summaries
    )
    _fill_grid(table, _with_empty_body(rows, 4))


def _fill_evidence_summaries(table: Table, draft: ReportDraft) -> None:
    rows = [("증빙 유형", "제출", "통과", "오류", "누락", "주요 내용")]
    rows.extend(
        (
            summary.evidence_type_name,
            _count(summary.submitted_count),
            _count(summary.passed_count),
            _count(summary.error_count),
            _count(summary.missing_count),
            summary.major_error,
        )
        for summary in draft.evidence_validation_summaries
    )
    _fill_grid(table, _with_empty_body(rows, 6))


def _fill_tax_rows(table: Table, draft: ReportDraft) -> None:
    rows = [("항목", "세금계산서 금액", "영수증 기재액", "부가세", "일치 여부")]
    rows.extend(
        (
            row.item_name,
            _money(row.document_supply_amount),
            _money(row.execution_supply_amount),
            _money(row.vat_amount),
            row.difference_label,
        )
        for row in draft.tax_settlement_rows
    )
    _fill_grid(table, _with_empty_body(rows, 5))


def _fill_item_reviews(table: Table, draft: ReportDraft) -> None:
    rows = [("No.", "집행 항목", "집행액", "판정", "요약 사유")]
    rows.extend(
        (
            str(review.no),
            review.item_name,
            _money(review.amount),
            review.decision_label,
            review.summary_reason,
        )
        for review in draft.item_reviews
    )
    _fill_grid(table, _with_empty_body(rows, 5))


def _fill_issue_detail_blocks(tables: Sequence[Table], issues: Sequence[IssueDetailDraft]) -> None:
    for table, issue in zip(tables, list(issues)[: len(tables)]):
        if issue.issue_type == "inappropriate":
            rows = [
                ("집행 금액", issue.amount_label),
                ("집행 경위", issue.problem or issue.agent_conclusion),
                ("판정 결론", "부적정"),
                ("법령 근거", issue.legal_basis),
                ("조치 사항", issue.required_action or issue.agent_conclusion),
            ]
        else:
            rows = [
                ("집행 금액", issue.amount_label),
                ("확인된 문제", issue.problem or issue.agent_conclusion),
                ("법령 근거", issue.legal_basis),
                ("필요 조치", issue.required_action or issue.agent_conclusion),
            ]
        _ensure_rows(table, len(rows))
        for row_index, row in enumerate(rows):
            _set_row(table, row_index, row)

    for table in tables[len(issues) :]:
        _fill_empty_issue_block(table)


def _ensure_issue_detail_blocks(
    document: DocumentObject,
    base_tables: Sequence[Table],
    issue_count: int,
) -> list[Table]:
    """6번 상세 블록 수를 이슈 개수에 맞춰 조정합니다."""

    desired_count = max(issue_count, 1)
    tables = list(base_tables)

    for index in range(len(tables) - 1, desired_count - 1, -1):
        heading = _find_paragraph_starting_with(document, f"6.{index + 1}")
        if heading is not None:
            _remove_element(heading._p)
        _remove_element(tables[index]._tbl)
    tables = tables[:desired_count]

    reference = _find_paragraph_starting_with(document, "7.")
    if reference is None:
        reference = document.add_paragraph()
    while len(tables) < desired_count:
        block_no = len(tables) + 1
        heading = document.add_paragraph()
        reference._p.addprevious(heading._p)
        _set_paragraph_text(heading, f"6.{block_no}  {EMPTY_LABEL}")

        table_xml = deepcopy(tables[-1]._tbl)
        reference._p.addprevious(table_xml)
        tables.append(Table(table_xml, document))

    return tables


def _fill_issue_headings(document: DocumentObject, issues: Sequence[IssueDetailDraft]) -> None:
    for index in range(max(len(issues), 1)):
        prefix = f"6.{index + 1}"
        paragraph = _find_paragraph_starting_with(document, prefix)
        if paragraph is None:
            continue
        if index >= len(issues):
            _set_paragraph_text(paragraph, f"{prefix}  {EMPTY_LABEL}")
            continue

        issue = issues[index]
        issue_label = "부적정" if issue.issue_type == "inappropriate" else "검토 필요"
        _set_paragraph_text(paragraph, f"{prefix}  {issue_label} - No.{issue.no}  {issue.title}")


def _fill_empty_issue_block(table: Table) -> None:
    rows = [
        ("집행 금액", EMPTY_LABEL),
        ("확인된 문제", EMPTY_LABEL),
        ("법령 근거", EMPTY_LABEL),
        ("필요 조치", EMPTY_LABEL),
    ]
    _ensure_rows(table, len(rows))
    for row_index, row in enumerate(rows):
        _set_row(table, row_index, row)


def _fill_supplement_actions(table: Table, draft: ReportDraft) -> None:
    actions = list(draft.supplement_actions)
    rows = [("No.", "보완 항목", "실행 내용", "완료 기한", "담당자")]
    rows.extend(
        (
            str(action.no),
            action.title,
            action.action,
            action.due_date_label,
            action.assignee,
        )
        for action in actions
    )
    _fill_grid(table, _with_empty_body(rows, 5))


def _fill_overall_opinion_paragraphs(document: DocumentObject, draft: ReportDraft) -> None:
    """v2 샘플의 문단형 종합 의견 섹션을 교체합니다."""

    heading_index = _find_paragraph_index(document, "8. 종합 의견")
    if heading_index is None:
        document.add_paragraph("8. 종합 의견")
        document.add_paragraph(draft.overall_opinion or EMPTY_LABEL)
        document.add_paragraph(f"검 토 자 :   {draft.reviewer_label}          (서명)  ______________________")
        document.add_paragraph(f"검토 일자 :   {draft.written_date_label}")
        document.add_paragraph(draft.department_label)
        return

    opinion_written = False
    for paragraph in document.paragraphs[heading_index + 1 :]:
        text = paragraph.text.strip()
        if text.startswith("검 토 자"):
            _set_paragraph_text(paragraph, f"검 토 자 :   {draft.reviewer_label}          (서명)  ______________________")
            opinion_written = True
            continue
        if text.startswith("검토 일자"):
            _set_paragraph_text(paragraph, f"검토 일자 :   {draft.written_date_label}")
            continue
        if opinion_written:
            if text and "안전관리팀" in text:
                _set_paragraph_text(paragraph, draft.department_label)
            continue
        if not opinion_written and text:
            _set_paragraph_text(paragraph, draft.overall_opinion or EMPTY_LABEL)
            opinion_written = True
        elif not opinion_written:
            continue
        else:
            _set_paragraph_text(paragraph, "")


def _fill_grid(table: Table, rows: Iterable[Sequence[str]]) -> None:
    materialized = [tuple(row) for row in rows]
    _ensure_rows(table, len(materialized))
    for row_index, row in enumerate(materialized):
        _set_row(table, row_index, row)


def _with_empty_body(rows: list[Sequence[str]], column_count: int) -> list[Sequence[str]]:
    if len(rows) == 1:
        rows.append((EMPTY_LABEL,) * column_count)
    return rows


def _ensure_rows(table: Table, count: int) -> None:
    while len(table.rows) < count:
        # 새 행도 음영과 테두리를 유지하도록 직전 행을 복제합니다.
        if table.rows:
            table._tbl.append(deepcopy(table.rows[-1]._tr))
        else:
            table.add_row()
    while len(table.rows) > count:
        table._tbl.remove(table.rows[-1]._tr)


def _set_row(table: Table, row_index: int, values: Sequence[str]) -> None:
    for column_index, value in enumerate(values[: len(table.columns)]):
        _set_cell(table.cell(row_index, column_index), value)
    for column_index in range(len(values), len(table.columns)):
        _set_cell(table.cell(row_index, column_index), "")


def _set_cell(cell: _Cell, value: object) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(value) if value not in (None, "") else EMPTY_LABEL)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(8.5)


def _find_paragraph_starting_with(document: DocumentObject, prefix: str):
    return next((paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith(prefix)), None)


def _find_paragraph_index(document: DocumentObject, text: str) -> int | None:
    return next((index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == text), None)


def _set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(9)


def _department_full_label(draft: ReportDraft) -> str:
    company = draft.basic_info.get("시공사", "").strip()
    department = draft.department_label.strip()
    if company and department and company not in department:
        return f"{company} {department}"
    return department or company or EMPTY_LABEL


def _remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _money(value: Decimal | int | str) -> str:
    return f"{Decimal(str(value)):,.0f}원"


def _count(value: int) -> str:
    return "-" if value == 0 else f"{value}건"


def load_report_draft(path: str | Path) -> ReportDraft:
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    return ReportDraft.model_validate(data)


def main() -> None:
    parser = argparse.ArgumentParser(description="Render ReportDraft JSON to DOCX.")
    parser.add_argument("input_json", help="Path to ReportDraft JSON.")
    parser.add_argument("output_docx", help="Path to write the rendered DOCX.")
    parser.add_argument(
        "--template",
        default=str(DEFAULT_TEMPLATE_PATH),
        help="Optional DOCX template path. Defaults to the analyzed report template.",
    )
    args = parser.parse_args()

    draft = load_report_draft(args.input_json)
    output = render_report_draft_to_docx(draft, args.output_docx, template_path=args.template)
    print(output)


if __name__ == "__main__":
    main()
